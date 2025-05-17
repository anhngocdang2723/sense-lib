import os
import logging
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional, Union
from langchain_text_splitters import RecursiveCharacterTextSplitter
import re
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from fastapi import HTTPException, UploadFile
from app.models import (
    User, Document, DocumentStatus, DocumentAccessLevel, FileType, Category, UserRole,
    DocumentChapter, DocumentSection, DocumentAudio, DocumentQA, ReadingProgress,
    DocumentAudioStatus, ReadingProgressType, ReadingProgressStatus, Voice
)
from app.models.language import Language
from app.schemas.document import DocumentCreate
from app.core.config import settings
from app.services.vector import VectorStore
from app.core.exceptions import (
    DuplicateFileError, InvalidFileError, FileProcessingError,
    VectorizationError, ValidationError, DatabaseError
)
from app.services.summary_service import SummaryService
from app.services.audio_service import AudioService

# Configure logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles text document processing including:
    - Text extraction
    - Document chunking with configurable overlap
    - Metadata extraction and management
    - Chapter and section detection
    - QA pair generation
    """
    
    SUPPORTED_EXTENSIONS = {
        'text': ['txt', 'pdf', 'docx', 'doc', 'rtf', 'odt', 'html', 'xml', 'json', 'csv', 'xls', 'xlsx', 'ppt', 'pptx'],
    }
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 200,
        upload_dir: str = settings.UPLOAD_DIR,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.upload_dir = upload_dir
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_dir, exist_ok=True)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )
    
    def process_file(self, file_path: str, db: Optional[Session] = None) -> Tuple[List[str], List[Dict[str, Any]]]:
        """
        Process a text file and return its chunks and associated metadata.
        
        Args:
            file_path: Path to the file to process
            db: Optional database session for storing chapters and sections
            
        Returns:
            Tuple containing:
                - List of text chunks
                - List of metadata dictionaries for each chunk
        """
        try:
            # Extract file extension
            file_ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            # Validate file type
            if not self._is_supported_file_type(file_ext):
                logger.error(f"Unsupported file type: {file_ext}")
                return [], []
            
            # Generate file hash for identification
            file_hash = self._generate_file_hash(file_path)
            
            # Extract basic metadata
            base_metadata = self._extract_metadata(file_path, file_ext)
            base_metadata["file_hash"] = file_hash
            
            # Extract text content
            text = self._extract_text(file_path)
            if not text:
                logger.error(f"Failed to extract text from {file_path}")
                return [], []
            
            # Detect chapters and sections
            chapters, sections = self._detect_structure(text)
            
            # Store chapters and sections in database if session provided
            if db:
                self._store_structure(db, file_hash, chapters, sections)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create metadata for each chunk
            metadata_list = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "content_type": "text",
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                })
                metadata_list.append(chunk_metadata)
            
            logger.info(f"Processed {file_path} into {len(chunks)} chunks")
            return chunks, metadata_list
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return [], []
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text content from file."""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                return self._clean_text(text)
            elif ext == '.pdf':
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                return self._clean_text(text)
            elif ext == '.docx':
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = '\n'.join([p.text for p in doc.paragraphs])
                return self._clean_text(text)
            else:
                logger.error(f'Unsupported file type: {ext}')
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove non-printable characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _is_supported_file_type(self, file_ext: str) -> bool:
        """Check if the file extension is supported."""
        return file_ext in self.SUPPORTED_EXTENSIONS['text']
    
    def _generate_file_hash(self, file_path: str) -> str:
        """Generate a hash for the file to use as a unique identifier."""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as file:
                # Read file in chunks to handle large files
                for chunk in iter(lambda: file.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Error generating file hash: {str(e)}")
            # Return a timestamp-based fallback hash
            return f"hash_{int(datetime.now().timestamp())}"
    
    def _extract_metadata(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """Extract metadata from the file."""
        try:
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            creation_time = os.path.getctime(file_path)
            modification_time = os.path.getmtime(file_path)
            
            metadata = {
                "file_name": file_name,
                "file_type": file_ext,
                "file_size": file_size,
                "creation_date": datetime.fromtimestamp(creation_time).isoformat(),
                "modification_date": datetime.fromtimestamp(modification_time).isoformat(),
                "source_path": file_path,
                "document_type": "text"
            }
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {
                "file_name": os.path.basename(file_path),
                "file_type": file_ext,
                "error": str(e)
            }
    
    def _detect_structure(self, text: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Detect chapters and sections in the text.
        
        Args:
            text: The text content to analyze
            
        Returns:
            Tuple containing:
                - List of chapter dictionaries
                - List of section dictionaries
        """
        chapters = []
        sections = []
        
        # Split text into lines
        lines = text.split('\n')
        
        current_chapter = None
        current_section = None
        chapter_number = 0
        section_number = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Check for chapter markers (e.g., "Chapter X" or "# Chapter X")
            chapter_match = re.match(r'^(?:#\s*)?(?:Chapter|CHAPTER)\s+(\d+)[:.]?\s*(.+)$', line, re.IGNORECASE)
            if chapter_match:
                chapter_number += 1
                current_chapter = {
                    "title": chapter_match.group(2).strip(),
                    "chapter_number": chapter_number,
                    "start_position": i / len(lines) * 100,  # Percentage through document
                    "end_position": None
                }
                chapters.append(current_chapter)
                section_number = 0
                continue
            
            # Check for section markers (e.g., "Section X" or "## Section X")
            section_match = re.match(r'^(?:##\s*)?(?:Section|SECTION)\s+(\d+)[:.]?\s*(.+)$', line, re.IGNORECASE)
            if section_match:
                section_number += 1
                current_section = {
                    "title": section_match.group(2).strip(),
                    "section_number": section_number,
                    "start_position": i / len(lines) * 100,
                    "end_position": None,
                    "chapter_number": chapter_number if current_chapter else None
                }
                sections.append(current_section)
                continue
        
        # Update end positions
        for i in range(len(chapters)):
            if i < len(chapters) - 1:
                chapters[i]["end_position"] = chapters[i + 1]["start_position"]
            else:
                chapters[i]["end_position"] = 100
        
        for i in range(len(sections)):
            if i < len(sections) - 1:
                sections[i]["end_position"] = sections[i + 1]["start_position"]
            else:
                sections[i]["end_position"] = 100
        
        return chapters, sections
    
    def _store_structure(
        self,
        db: Session,
        file_hash: str,
        chapters: List[Dict[str, Any]],
        sections: List[Dict[str, Any]]
    ) -> None:
        """
        Store detected chapters and sections in the database.
        
        Args:
            db: Database session
            file_hash: Hash of the document file
            chapters: List of chapter dictionaries
            sections: List of section dictionaries
        """
        try:
            # Get document ID from file hash
            document = db.query(Document).filter(Document.file_hash == file_hash).first()
            if not document:
                logger.error(f"Document not found for file hash: {file_hash}")
                return
            
            # Store chapters
            for chapter_data in chapters:
                chapter = DocumentChapter(
                    document_id=document.id,
                    title=chapter_data["title"],
                    chapter_number=chapter_data["chapter_number"],
                    start_position=chapter_data["start_position"],
                    end_position=chapter_data["end_position"]
                )
                db.add(chapter)
            
            db.commit()
            
            # Store sections
            for section_data in sections:
                section = DocumentSection(
                    document_id=document.id,
                    title=section_data["title"],
                    section_number=section_data["section_number"],
                    start_position=section_data["start_position"],
                    end_position=section_data["end_position"]
                )
                db.add(section)
            
            db.commit()
            
        except Exception as e:
            logger.error(f"Error storing document structure: {str(e)}")
            db.rollback()
    
    def generate_qa_pairs(self, text: str, context: Optional[str] = None) -> List[Dict[str, str]]:
        """
        Generate question-answer pairs from text.
        
        Args:
            text: The text to generate QA pairs from
            context: Optional context for the QA pairs
            
        Returns:
            List of dictionaries containing question-answer pairs
        """
        # TODO: Implement QA pair generation using LLM
        return []

    @staticmethod
    async def get_file_hash(file: UploadFile) -> str:
        """Generate a hash for the uploaded file."""
        try:
            hasher = hashlib.md5()
            content = await file.read()
            hasher.update(content)
            await file.seek(0)  # Reset file pointer
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Error generating file hash: {str(e)}")
            return f"hash_{int(datetime.now().timestamp())}"

class DocumentService:
    def __init__(self):
        self.summary_service = SummaryService()
        self.audio_service = AudioService()
    
    @staticmethod
    def validate_isbn(isbn: Optional[str]) -> bool:
        """Validate ISBN format"""
        if isbn is None:
            return True
        # Check format: ISBN-10, ISBN-13, or ISBN-13 with hyphens
        return bool(re.match(r'^(?:[0-9]{10}|[0-9]{13}|[0-9]{3}-[0-9]{1,5}-[0-9]{1,7}-[0-9]{1,6}-[0-9])$', isbn))

    @staticmethod
    def validate_publication_year(year: Optional[int]) -> bool:
        """Validate publication year"""
        if year is None:
            return True
        current_year = datetime.now().year
        return 1800 <= year <= current_year

    @staticmethod
    def validate_file_size(file_size: int) -> bool:
        """Validate file size"""
        return file_size > 0

    @staticmethod
    def validate_version(version: str) -> bool:
        """Validate version format (e.g., 1.0, 1.0.0)"""
        return bool(re.match(r'^[0-9]+\.[0-9]+(\.[0-9]+)?$', version))

    @staticmethod
    async def validate_document_data(
        db: Session,
        data: DocumentCreate,
        file: UploadFile,
        current_user: User
    ) -> Tuple[bool, str]:
        """Validate all document data before processing"""
        logger.info("Starting document data validation")
        
        # Check required fields
        if not data.title:
            logger.error("Validation failed: Title is required")
            raise ValidationError(
                "Title is required",
                data={"field": "title"}
            )
        
        # Validate ISBN
        logger.info(f"Validating ISBN: {data.isbn}")
        if not DocumentService.validate_isbn(data.isbn):
            logger.error("Validation failed: Invalid ISBN format")
            raise ValidationError(
                "Invalid ISBN format. Must be in format: 1234567890, 9781234567890, or 978-604-1-00614-6",
                data={"field": "isbn", "value": data.isbn}
            )
        
        # Validate publication year
        logger.info(f"Validating publication year: {data.publication_year}")
        if not DocumentService.validate_publication_year(data.publication_year):
            logger.error("Validation failed: Invalid publication year")
            raise ValidationError(
                "Invalid publication year. Must be between 1800 and current year",
                data={"field": "publication_year", "value": data.publication_year}
            )
        
        # Validate file size
        logger.info(f"Validating file size: {file.size} bytes")
        if not DocumentService.validate_file_size(file.size):
            logger.error("Validation failed: Invalid file size")
            raise InvalidFileError(
                "Invalid file size. Must be greater than 0",
                data={"file_size": file.size}
            )
        
        # Validate version
        logger.info(f"Validating version: {data.version}")
        if not DocumentService.validate_version(data.version):
            logger.error("Validation failed: Invalid version format")
            raise ValidationError(
                "Invalid version format. Must be in format x.y or x.y.z",
                data={"field": "version", "value": data.version}
            )
        
        # Check if language exists
        logger.info(f"Checking language existence: {data.language}")
        if data.language:
            language = db.query(Language).filter(Language.code == data.language).first()
            if not language:
                logger.error(f"Validation failed: Language {data.language} does not exist")
                raise ValidationError(
                    f"Language with code {data.language} does not exist",
                    data={"field": "language", "value": data.language}
                )
        
        # Check if category exists
        logger.info(f"Checking category existence: {data.category_id}")
        if data.category_id:
            category = db.query(Category).filter(Category.id == data.category_id).first()
            if not category:
                logger.error(f"Validation failed: Category {data.category_id} does not exist")
                raise ValidationError(
                    f"Category with ID {data.category_id} does not exist",
                    data={"field": "category_id", "value": str(data.category_id)}
                )
        
        logger.info("Document data validation completed successfully")
        return True, ""

    @staticmethod
    def _check_duplicate_file(file_hash: str, db: Session) -> None:
        """Check if a file with the same hash already exists"""
        existing_doc = db.query(Document).filter(Document.file_hash == file_hash).first()
        if existing_doc:
            raise DuplicateFileError(file_hash=file_hash, document_id=str(existing_doc.id))

    @staticmethod
    async def process_and_save_document(
        db: Session,
        data: DocumentCreate,
        file: UploadFile,
        current_user: User
    ) -> Document:
        """Process and save document with all validations"""
        logger.info("Starting document processing and saving")
        
        try:
            # Validate all data first
            logger.info("Validating document data")
            await DocumentService.validate_document_data(db, data, file, current_user)

            # Calculate file hash
            logger.info("Calculating file hash")
            file_hash = await DocumentProcessor.get_file_hash(file)
            logger.info(f"File hash calculated: {file_hash}")
            
            # Check for duplicate file
            logger.info("Checking for duplicate file")
            DocumentService._check_duplicate_file(file_hash, db)

            # Save file to disk
            logger.info("Saving file to disk")
            timestamp = int(datetime.utcnow().timestamp())
            safe_filename = f"{timestamp}_{file.filename}"
            file_path = os.path.join(settings.UPLOAD_DIR, safe_filename)
            logger.info(f"File will be saved as: {safe_filename}")
            
            try:
                # Read file content
                logger.info("Reading file content")
                content = await file.read()
                with open(file_path, "wb") as buffer:
                    buffer.write(content)
                logger.info("File saved successfully")
            except Exception as e:
                logger.error(f"Error saving file: {str(e)}")
                raise FileProcessingError(
                    "Failed to save file to disk",
                    data={"error": str(e)}
                )

            # Get file type
            logger.info("Getting file type")
            file_ext = os.path.splitext(file.filename)[1].lower().replace('.', '')
            file_type = db.query(FileType).filter(FileType.extension == file_ext).first()
            if not file_type:
                logger.error(f"Unsupported file type: {file_ext}")
                raise InvalidFileError(
                    f"Unsupported file type: {file_ext}",
                    data={"file_type": file_ext}
                )
            logger.info(f"File type identified: {file_ext}")

            try:
                # Create document instance
                logger.info("Creating document instance")
                document_data = data.model_dump()
                document_data.update({
                    "file_name": safe_filename,
                    "file_hash": file_hash,
                    "file_size": len(content),
                    "file_type": file_type.id,
                    "added_by": current_user.id,
                    "status": DocumentStatus.PENDING
                })
                
                db_document = Document(**document_data)
                db.add(db_document)
                db.commit()
                db.refresh(db_document)
                logger.info(f"Document instance created with ID: {db_document.id}")
            except Exception as e:
                logger.error(f"Database error: {str(e)}")
                raise DatabaseError(
                    "Failed to create document record",
                    data={"error": str(e)}
                )

            # Process document using DocumentProcessor
            logger.info("Starting document processing")
            processor = DocumentProcessor()
            chunks, metadata_list = processor.process_file(file_path, db)
            
            if not chunks:
                logger.error("Document processing failed - no chunks generated")
                db_document.status = DocumentStatus.REJECTED
                db.commit()
                raise FileProcessingError(
                    "Failed to process document content",
                    data={"document_id": str(db_document.id)}
                )
            logger.info(f"Document processed into {len(chunks)} chunks")

            # Generate summary from the first chunk (or combine chunks if needed)
            logger.info("Generating document summary")
            summary_service = SummaryService()
            summary = await summary_service.generate_summary(" ".join(chunks))
            db_document.ai_summary = summary
            
            # Comment out vector store operations
            """
            # Add to vector store
            try:
                logger.info("Initializing vector store")
                vector_store = VectorStore(
                    qdrant_url=settings.QDRANT_URL,
                    qdrant_api_key=settings.QDRANT_API_KEY
                )
                
                logger.info("Storing document chunks in vector store")
                success = vector_store.store_documents(chunks, metadata_list)
                if not success:
                    logger.error("Failed to store document in vector store")
                    raise VectorizationError(
                        "Failed to store document in vector store",
                        data={"document_id": str(db_document.id)}
                    )
                
                logger.info("Document successfully stored in vector store")
                db_document.status = DocumentStatus.AVAILABLE
                db.commit()
                logger.info("Document status updated to AVAILABLE")
            except Exception as e:
                logger.error(f"Vector store operation failed: {str(e)}", exc_info=True)
                db_document.status = DocumentStatus.REJECTED
                db.commit()
                raise VectorizationError(
                    "Document saved but vectorization failed",
                    data={
                        "document_id": str(db_document.id),
                        "error": str(e)
                    }
                )
            """

            # Generate audio from summary
            logger.info("Generating audio from summary")
            audio_service = AudioService()
            
            # Create introduction text with book title
            intro_text = f"Sau đây là bản tóm tắt của {db_document.title}"
            # Remove author check since it's not in the model
            intro_text += ". "
            
            # Combine introduction with summary
            full_text = intro_text + summary
            
            # Get default voice for the language
            default_voice = db.query(Voice).filter(
                Voice.language == data.language,
                Voice.is_active == True
            ).first()
            
            if not default_voice:
                logger.warning(f"No default voice found for language {data.language}, using system default")
                current_time = datetime.utcnow()
                default_voice = Voice(
                    id=settings.DEFAULT_VOICE_ID,
                    name="Default Voice",
                    language=data.language,
                    provider="gTTS",
                    is_active=True,
                    created_at=current_time,
                    updated_at=current_time,
                    gender=None  # Optional field
                )
                try:
                    db.add(default_voice)
                    db.commit()
                    logger.info(f"Created default voice for language {data.language}")
                except Exception as e:
                    logger.error(f"Error creating default voice: {str(e)}")
                    db.rollback()
                    raise DatabaseError(
                        "Failed to create default voice",
                        data={"error": str(e)}
                    )
            
            # Generate audio using original filename
            original_filename = os.path.splitext(db_document.file_name)[0]  # Remove extension
            audio_result = await audio_service.generate_audio(
                text=full_text,
                language=data.language,
                filename=original_filename
            )
            
            # Create audio record with all required fields
            current_time = datetime.utcnow()
            document_audio = DocumentAudio(
                document_id=db_document.id,
                language=data.language,
                voice_id=default_voice.id,
                file_url=audio_result["file_url"],
                duration_seconds=audio_result["duration_seconds"],
                file_size=audio_result["file_size"],
                status=DocumentAudioStatus.COMPLETED,
                created_at=current_time,
                updated_at=current_time
            )
            
            try:
                db.add(document_audio)
                db.commit()
                logger.info(f"Audio record created successfully for document {db_document.id}")
            except Exception as e:
                logger.error(f"Error creating audio record: {str(e)}")
                db.rollback()
                raise DatabaseError(
                    "Failed to create audio record",
                    data={"error": str(e)}
                )

            # Update document status and commit
            db_document.status = DocumentStatus.AVAILABLE
            db.commit()
            
            logger.info("Document processing completed successfully")
            return db_document
            
        except Exception as e:
            logger.error(f"Error in document processing: {str(e)}")
            if 'db_document' in locals():
                db_document.status = DocumentStatus.REJECTED
                db.commit()
            raise

    def __del__(self):
        self.summary_service = None
        self.audio_service = None 